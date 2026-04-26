"""read_docx — extract text from a Microsoft Word .docx file."""
from pathlib import Path
from tools.base import BaseTool, InputSchema, ToolProperty, ToolWeight
from logger import get_logger

logger = get_logger(__name__)

_INLINE_CAP = 40_000


class ReadDocxTool(BaseTool):
    name = "read_docx"
    description = (
        "Extract text from a Microsoft Word .docx file. "
        "Includes body paragraphs and tables as plain text. "
        "Set artifact_key to store the extracted text as a named artifact."
    )
    weight = ToolWeight.MODERATE

    @property
    def input_schema(self) -> InputSchema:
        return InputSchema(
            properties={
                "path": ToolProperty(
                    type="string",
                    description="Path to the .docx file",
                ),
                "include_headers_footers": ToolProperty(
                    type="boolean",
                    description="Include header and footer text (default false)",
                ),
                "artifact_key": ToolProperty(
                    type="string",
                    description="Store extracted text under this artifact key for later reading",
                ),
            },
            required=["path"],
        )

    def execute(self, tool_input: dict) -> str:
        try:
            import docx
        except ImportError:
            return "Error: python-docx is not installed. Run: pip install python-docx"

        path = Path(tool_input["path"])
        if not path.exists():
            return f"Error: file not found: {path}"
        if not path.is_file():
            return f"Error: not a file: {path}"

        include_hf = tool_input.get("include_headers_footers", False)
        artifact_key = (tool_input.get("artifact_key") or "").strip()

        try:
            doc = docx.Document(str(path))
        except Exception as e:
            return f"Error: could not open '{path.name}' as a DOCX file: {e}"

        parts = []

        # Body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tables
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                parts.append(f"\n[Table {i + 1}]")
                parts.extend(rows)

        # Headers / footers
        if include_hf:
            for section in doc.sections:
                for header in [section.header, section.footer]:
                    if header is not None:
                        for para in header.paragraphs:
                            text = para.text.strip()
                            if text:
                                parts.append(f"[Header/Footer] {text}")

        full_text = "\n".join(parts)

        table_count = len(doc.tables)
        para_count = len([p for p in doc.paragraphs if p.text.strip()])

        # Store artifact if requested
        if artifact_key:
            try:
                from runtime.artifact_store import get_artifact_store
                store = get_artifact_store()
                store.set(artifact_key, full_text, kind="url_content", source=str(path))
                logger.info(f"  read_docx: stored {len(full_text)} chars as artifact '{artifact_key}'")
            except Exception as e:
                logger.warning(f"  read_docx: artifact store unavailable: {e}")

        header = (
            f"DOCX: {path}\n"
            f"Paragraphs: {para_count}  |  Tables: {table_count}\n"
        )

        if len(full_text) > _INLINE_CAP:
            truncated = full_text[:_INLINE_CAP]
            note = f"\n[truncated at {_INLINE_CAP:,} chars — {len(full_text):,} total. "
            if artifact_key:
                note += f"Use get_artifact '{artifact_key}' to read the rest.]"
            else:
                note += f"Re-run with artifact_key='docx_content' to store the full text.]"
            return header + "\n" + truncated + note
        else:
            return header + "\n" + full_text
