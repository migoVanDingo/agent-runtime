"""document_info — metadata and statistics for PDF or DOCX files."""
from pathlib import Path
from tools.base import BaseTool, InputSchema, ToolProperty, ToolWeight
from logger import get_logger

logger = get_logger(__name__)


def _fmt_size(size_bytes: int) -> str:
    if size_bytes >= 1_000_000:
        return f"{size_bytes / 1_000_000:.1f} MB"
    if size_bytes >= 1_000:
        return f"{size_bytes / 1_000:.1f} kB"
    return f"{size_bytes} B"


def _fmt_ts(ts) -> str:
    """Format a datetime or string timestamp."""
    if ts is None:
        return ""
    if hasattr(ts, "strftime"):
        return ts.strftime("%Y-%m-%d %H:%M")
    return str(ts)


class DocumentInfoTool(BaseTool):
    name = "document_info"
    description = (
        "Show metadata and statistics for a PDF or DOCX file — "
        "title, author, page/paragraph count, file size, and dates. "
        "Does not extract full text; use read_pdf or read_docx for that."
    )
    weight = ToolWeight.LIGHTWEIGHT

    @property
    def input_schema(self) -> InputSchema:
        return InputSchema(
            properties={
                "path": ToolProperty(
                    type="string",
                    description="Path to a PDF or DOCX file",
                ),
            },
            required=["path"],
        )

    def execute(self, tool_input: dict) -> str:
        path = Path(tool_input["path"])
        if not path.exists():
            return f"Error: file not found: {path}"
        if not path.is_file():
            return f"Error: not a file: {path}"

        ext = path.suffix.lower()
        size_str = _fmt_size(path.stat().st_size)

        if ext == ".pdf":
            return self._pdf_info(path, size_str)
        elif ext in (".docx", ".doc"):
            return self._docx_info(path, size_str)
        else:
            return f"Error: unsupported format '{ext}'. Supported: .pdf, .docx"

    def _pdf_info(self, path: Path, size_str: str) -> str:
        try:
            import pypdf
        except ImportError:
            return "Error: pypdf is not installed. Run: pip install pypdf"

        try:
            reader = pypdf.PdfReader(str(path))
        except Exception as e:
            return f"Error: could not open PDF: {e}"

        pages = len(reader.pages)
        encrypted = reader.is_encrypted
        meta = {}
        if reader.metadata:
            meta = dict(reader.metadata)

        def _m(key: str) -> str:
            val = meta.get(f"/{key}", meta.get(key, ""))
            return str(val).strip() if val else ""

        lines = [
            f"Format: PDF",
            f"File:   {path.name} ({size_str})",
            f"Pages:  {pages}",
            f"Encrypted: {'Yes' if encrypted else 'No'}",
        ]
        if _m("Title"):
            lines.append(f"Title:  {_m('Title')}")
        if _m("Author"):
            lines.append(f"Author: {_m('Author')}")
        if _m("Creator"):
            lines.append(f"Creator: {_m('Creator')}")
        if _m("Producer"):
            lines.append(f"Producer: {_m('Producer')}")
        if _m("CreationDate"):
            lines.append(f"Created: {_m('CreationDate')}")
        if _m("ModDate"):
            lines.append(f"Modified: {_m('ModDate')}")

        return "\n".join(lines)

    def _docx_info(self, path: Path, size_str: str) -> str:
        try:
            import docx
        except ImportError:
            return "Error: python-docx is not installed. Run: pip install python-docx"

        try:
            doc = docx.Document(str(path))
        except Exception as e:
            return f"Error: could not open DOCX: {e}"

        props = doc.core_properties
        para_count = len([p for p in doc.paragraphs if p.text.strip()])
        table_count = len(doc.tables)

        # Estimate word count
        word_count = sum(
            len(p.text.split()) for p in doc.paragraphs if p.text.strip()
        )

        lines = [
            f"Format: DOCX",
            f"File:   {path.name} ({size_str})",
            f"Paragraphs: {para_count}",
            f"Tables: {table_count}",
            f"Words (estimated): ~{word_count:,}",
        ]
        if props.title:
            lines.append(f"Title:  {props.title}")
        if props.author:
            lines.append(f"Author: {props.author}")
        if props.last_modified_by:
            lines.append(f"Last Modified By: {props.last_modified_by}")
        if props.created:
            lines.append(f"Created: {_fmt_ts(props.created)}")
        if props.modified:
            lines.append(f"Modified: {_fmt_ts(props.modified)}")
        if props.description:
            lines.append(f"Description: {props.description[:200]}")

        return "\n".join(lines)
